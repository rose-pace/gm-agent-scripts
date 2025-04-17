from typing import Any, Dict, List, Optional
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from rich.console import Console
from rich.table import Table as RichTable
from statblock_validator import StatBlockValidator
from validators.ability_validators import calculate_proficiency_bonus

# Import all parsers
from parsers.base_parser import BaseParser
from parsers.core_stats_parser import CoreStatsParser
from parsers.abilities_parser import AbilitiesParser
from parsers.actions_parser import ActionsParser
from parsers.legendary_actions_parser import LegendaryActionsParser
from parsers.lair_actions_parser import LairActionsParser
from parsers.regional_effects_parser import RegionalEffectsParser
from parsers.description_parser import DescriptionParser
from parsers.spellcasting_parser import SpellcastingParser
from parsers.damage_type_parser import DamageTypeParser

class DocxStatBlockConverter(BaseParser):
    def __init__(self, collection: str = None, tags: List[str] = None):
        """Initialize converter with validation rules."""
        self.collection = collection
        self.tags = tags or []
        self.current_creature = {}
        self.console = Console()
        
    @staticmethod
    def _is_run_bold(run) -> bool:
        """
        Check if a run is bold using multiple methods.
        - Checks direct bold property
        - Checks for 'Strong' style
        - Checks font.bold property
        """
        return (
            run.bold or  # Direct bold property
            (hasattr(run, 'style') and run.style and 'Strong' in run.style.name) or  # Style name
            (hasattr(run, 'font') and run.font and run.font.bold)  # Font property
        )

    def extract_text_from_docx(self, docx_path: str) -> tuple[Dict[str, List[Paragraph]], List[Table]]:
        """
        Extract text from DOCX file while preserving formatting.
        Returns dictionary of sections with their paragraphs.
        """
        doc = Document(docx_path)
        sections = {}
        current_section = 'corestats'
        current_paragraphs = []

        for paragraph in doc.paragraphs:
            # Check if this is a main header
            if self._is_main_header(paragraph):
                sections['header'] = BaseParser.normalize_text(paragraph.text)
            # Check if this is a subheader
            elif self._is_subheader(paragraph):
                sections['subheader'] = BaseParser.normalize_text(paragraph.text)
            # Check if this is a section header
            elif self._is_section_header(paragraph):
                # Save previous section
                if current_paragraphs:
                    sections[current_section] = current_paragraphs
                # Start new section
                current_section = self._get_section_name(paragraph.text)
                current_paragraphs = []
            else:
                # Skip empty paragraphs
                if paragraph.text.strip():
                    current_paragraphs.append(paragraph)

        # Add last section
        if current_paragraphs:
            sections[current_section] = current_paragraphs

        return sections, doc.tables
    
    def _is_main_header(self, paragraph: Paragraph) -> bool:
        """
        Check if paragraph is a main header based on formatting.
        """
        return paragraph.style.name.lower() == 'heading 1'
    
    def _is_subheader(self, paragraph: Paragraph) -> bool:
        """
        Check if paragraph is a subheader based on formatting.
        """
        return paragraph.style.name.lower() == 'heading 2'

    def _is_section_header(self, paragraph: Paragraph) -> bool:
        """
        Check if paragraph is a section header based on formatting.
        """
        # Check for bold formatting or specific styles
        return paragraph.style.name.lower() == 'heading 3'

    def _get_section_name(self, text: str) -> str:
        """Convert section header text to section identifier."""
        clean_text = text.lower().strip().strip('.')
        return self.SECTION_MARKERS.get(clean_text, clean_text)

    def convert_docx_to_schema(self, docx_path: str) -> Dict:
        """Convert DOCX stat block to schema format."""
        sections, tables = self.extract_text_from_docx(docx_path)
        
        # Initialize with required fields
        self.current_creature = {
            'metadata': {
                'name': sections['header'],  # Name belongs in metadata
                'version': '1.0',
                'date_created': datetime.now().strftime('%Y-%m-%d'),
                'last_modified': datetime.now().strftime('%Y-%m-%d'),
                'source': Path(docx_path).name,
                'collection': self.collection,
                'tags': self.tags
            },
            'creature_info': {},
            'core_stats': {},
            'abilities': {},
            'proficiencies': {
                'saving_throws': [],
                'skills': [],
                'bonus': 2
            },
            'defenses': {
                'damage_resistances': [],
                'damage_immunities': [],
                'condition_immunities': [],
            },            
            'senses': {
                'passive_perception': 10,
                'special': []
            },
            'languages': {
                'spoken': []
            },
        }
        
        # Process main sections
        self._process_subheader(sections['subheader'])
        self._process_core_stats(sections.get('corestats', []))
        self._process_abilities(sections.get('corestats', []), tables)
        self._process_defenses(sections.get('corestats', []))
        self._process_traits(sections.get('traits', []))
        
        # Process actions
        standard_actions = self._process_actions(sections.get('actions', []), "standard")
        bonus_actions = self._process_actions(sections.get('bonus_actions', []), "bonus")
        reactions = self._process_actions(sections.get('reactions', []), "reaction")
        
        self.current_creature["actions"] = {
            'standard': standard_actions,
            'bonus_actions': bonus_actions if bonus_actions else None,
            'reactions': reactions if reactions else None
        }
        
        # Process optional sections
        if 'legendary_actions' in sections:
            self._process_legendary_actions(sections['legendary_actions'])
        if 'lair_actions' in sections:
            self._process_lair_actions(sections['lair_actions'])
        if 'regional_effects' in sections:
            self._process_regional_effects(sections['regional_effects'])
        if 'description' in sections:
            self._process_description(sections['description'])

        # Validate converted data
        self._validate_converted_data()
        
        return self.current_creature

    def _validate_converted_data(self) -> None:
        """
        Validate converted data using Pydantic model.
        """
        try:
            StatBlockValidator(**self.current_creature)
        except Exception as e:
            self.console.print("[red]Validation Error:[/red]")
            self.console.print(e)
            raise

    def output_conversion_report(self) -> None:
        """
        Generate a detailed report of the conversion process.
        """
        table = RichTable(title="Conversion Report")
        table.add_column("Section", style="cyan")
        table.add_column("Status", style="green")
        table.add_column("Notes", style="yellow")
        
        # Add rows for each processed section
        for section, status in self.conversion_status.items():
            table.add_row(
                section,
                "✓" if status['success'] else "✗",
                status.get('notes', '')
            )
        
        # Print to console
        self.console.print(table)

    # Constants and patterns
    SECTION_MARKERS = {
        'actions': 'actions',
        'legendary actions': 'legendary_actions',
        'lair actions': 'lair_actions',
        'regional effects': 'regional_effects',
        'traits': 'traits',
        'bonus actions': 'bonus_actions',
        'reactions': 'reactions',
        'description': 'description'
    }

    MELEE_ATTACK_PATTERN = r'^Melee (?:Weapon|Spell) Attack:\s*(?P<bonus>[+-]\d+) to hit, reach (?P<distance>\d+ ft\.)'
    RANGED_ATTACK_PATTERN = r'^Ranged (?:Weapon|Spell) Attack:\s*(?P<bonus>[+-]\d+) to hit, range (?P<distance>\d+/\d+ ft\.)'

    USAGE_PATTERNS = {
        'recharge': r'recharge (\d+)(?:-(\d+))?',
        'per_day': r'(\d+)/day',
        'per_short_rest': r'(\d+)/short rest',
        'per_long_rest': r'(\d+)/long rest',
        'costs': r'costs? (\d+)',
    }

    def _process_subheader(self, subheader: str) -> None:
        """Process header section containing size, type, alignment."""
        if not subheader:
            raise ValueError("Incomplete header section")
        
        subheader_data = CoreStatsParser.parse_subheader(subheader)
        self.current_creature['creature_info'].update(subheader_data)

    def _process_core_stats(self, paragraphs: List[Paragraph]) -> None:
        """Process core statistics section."""
        core_stats = {}

        for para in paragraphs:
            text = BaseParser.normalize_text(para.text)
            
            if text.startswith("Armor Class"):
                core_stats["armor_class"] = CoreStatsParser.parse_armor_class(text)
            elif text.startswith("Hit Points"):
                core_stats["hit_points"] = CoreStatsParser.parse_hit_points(text)
            elif text.startswith("Speed"):
                core_stats["speed"] = CoreStatsParser.parse_speed(text)
            elif text.startswith("Senses"):
                self.current_creature["senses"] = CoreStatsParser.parse_senses(text)
            elif text.startswith("Languages"):
                self.current_creature["languages"] = CoreStatsParser.parse_languages(text)
            elif text.startswith("Challenge"):
                cr = CoreStatsParser.parse_challenge_rating(text)
                self.current_creature['creature_info']['cr'] = cr
                self.current_creature['proficiencies']['bonus'] = calculate_proficiency_bonus(cr['rating'])

        self.current_creature['core_stats'] = core_stats

    def _process_abilities(self, paragraphs: List[Paragraph], tables: List[Table]) -> None:
        """Process abilities section including table-based formats."""
        for table in tables:
            if len(table.rows) >= 2 and len(table.rows[0].cells) >= 6:
                # Check if this is the ability score table
                header_cells = [cell.text.strip().upper() for cell in table.rows[0].cells]
                if all(ability in header_cells for ability in ['STR', 'DEX', 'CON', 'INT', 'WIS', 'CHA']):
                    abilities = {}
                    ability_values = table.rows[1].cells
                    for i, ability in enumerate(['strength', 'dexterity', 'constitution', 'intelligence', 'wisdom', 'charisma']):
                        abilities[ability] = AbilitiesParser.parse_ability_scores(ability_values[i].text)
                    self.current_creature["abilities"] = abilities
                    break

        if not self.current_creature['core_stats'].get('initiative'):
            # If initiative is not in core stats, calculate it from abilities
            if 'dexterity' in self.current_creature['abilities']:
                dexterity = self.current_creature['abilities']['dexterity']
                self.current_creature['core_stats']['initiative'] = {
                    'bonus': dexterity['modifier'],
                    'average': dexterity['modifier'] + 10  
                }
            else:
                self.current_creature['core_stats']['initiative'] = {
                    'bonus': 0,
                    'average': 10  
                }

        for para in paragraphs:
            text = BaseParser.normalize_text(para.text)
            if text.startswith("Saving Throws"):
                self.current_creature['proficiencies']["saving_throws"] = AbilitiesParser.parse_saving_throws(text)
            elif text.startswith("Skills"):
                self.current_creature['proficiencies']["skills"] = AbilitiesParser.parse_skills(text)

    def _process_actions(self, paragraphs: List[Paragraph], action_type: str = "standard") -> List[Dict]:
        """Process actions section."""
        actions = []
        current_action = None
        
        for para in paragraphs:
            if para.runs and self._is_run_bold(para.runs[0]):
                if current_action:
                    # Check if current action is spellcasting
                    if 'spellcasting' in current_action['name'].lower():
                        spellcasting = SpellcastingParser().parse_spellcasting_trait(
                            f'{current_action['name']} {current_action['description']}', 
                            self.current_creature['abilities']
                        )
                        if spellcasting:
                            self.current_creature['spellcasting'] = spellcasting
                    else:
                        actions.append(current_action)
                
                name, description = BaseParser.split_name_description(para.text)
                current_action = ActionsParser.parse_action(description, name)
            elif current_action:
                current_action["description"] += f"\n{BaseParser.normalize_text(para.text)}"
        
        # Handle last action
        if current_action:
            if 'spellcasting' in current_action['name'].lower():
                spellcasting = SpellcastingParser().parse_spellcasting_trait(
                    f'{current_action['name']} {current_action['description']}', 
                    self.current_creature['abilities']
                )
                if spellcasting:
                    self.current_creature['spellcasting'] = spellcasting
            else:
                actions.append(current_action)
        
        return actions

    def _process_legendary_actions(self, paragraphs: List[Paragraph]) -> None:
        """Process legendary actions section."""
        if not paragraphs:
            return
            
        header_text = BaseParser.normalize_text(paragraphs[0].text)
        action_texts = [BaseParser.normalize_text(p.text) for p in paragraphs[1:]]
        self.current_creature["legendary_actions"] = LegendaryActionsParser.parse_legendary_actions(
            header_text, action_texts
        )

    def _process_lair_actions(self, paragraphs: List[Paragraph]) -> None:
        """Process lair actions section."""
        if not paragraphs:
            return
            
        header_text = BaseParser.normalize_text(paragraphs[0].text)
        action_texts = [BaseParser.normalize_text(p.text) for p in paragraphs[1:]]
        self.current_creature["lair_actions"] = LairActionsParser.parse_lair_actions(
            header_text, action_texts
        )

    def _process_regional_effects(self, paragraphs: List[Paragraph]) -> None:
        """Process regional effects section."""
        if not paragraphs:
            return
            
        header_text = BaseParser.normalize_text(paragraphs[0].text)
        effect_texts = [BaseParser.normalize_text(p.text) for p in paragraphs[1:]]
        self.current_creature["regional_effects"] = RegionalEffectsParser.parse_regional_effects(
            header_text, effect_texts
        )

    def _process_defenses(self, paragraphs: List[Paragraph]) -> None:
        """Process damage and condition immunities/resistances."""
        for para in paragraphs:
            text = BaseParser.normalize_text(para.text)
            
            if text.startswith("Damage Resistances"):
                resistances = text.replace("Damage Resistances", "").strip()
                self.current_creature['defenses']["damage_resistances"] = DamageTypeParser.parse_damage_types(resistances)
            elif text.startswith("Damage Immunities"):
                immunities = text.replace("Damage Immunities", "").strip()
                self.current_creature['defenses']["damage_immunities"] = DamageTypeParser.parse_damage_types(immunities)
            elif text.startswith("Condition Immunities"):
                conditions = text.replace("Condition Immunities", "").strip()
                self.current_creature['defenses']["condition_immunities"] = [c.strip() for c in conditions.split(", ")]

    def _process_traits(self, paragraphs: List[Paragraph]) -> None:
        """Process traits section with spellcasting detection."""
        traits = []
        current_trait = None
        spellcasting_parser = SpellcastingParser()

        for para in paragraphs:
            if para.runs and self._is_run_bold(para.runs[0]):
                if current_trait:
                    # Check if current trait is spellcasting
                    if 'spellcasting' in current_trait['name'].lower():
                        spellcasting = spellcasting_parser.parse_spellcasting_trait(f'{current_trait['name']} {current_trait['description']}', self.current_creature['abilities'])
                        if spellcasting:
                            self.current_creature['spellcasting'] = spellcasting
                    else:
                        traits.append(current_trait)

                name, description = BaseParser.split_name_description(para.text)
                current_trait = {
                    'name': name,
                    'description': description
                }
            elif current_trait:
                current_trait['description'] += f'\n{BaseParser.normalize_text(para.text)}'

        # Handle last trait
        if current_trait:
            if 'spellcasting' in current_trait['name'].lower():
                spellcasting = spellcasting_parser.parse_spellcasting_trait(f'{current_trait['name']} {current_trait['description']}', self.current_creature['abilities'])
                if spellcasting:
                    self.current_creature['spellcasting'] = spellcasting
            else:
                traits.append(current_trait)

        self.current_creature['traits'] = traits if traits else None

    def _process_description(self, paragraphs: List[Paragraph]) -> None:
        """Process description section."""
        text = "\n".join([BaseParser.normalize_text(para.text) for para in paragraphs])
        self.current_creature['description'] = DescriptionParser.classify_text(text)
